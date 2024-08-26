import streamlit as st
import msal
import requests
import pandas as pd
import time
from io import BytesIO
from dotenv import load_dotenv
import os
from openai import OpenAI
import json
from docx import Document
from zipfile import ZipFile

load_dotenv()
API_KEY = os.getenv('API_KEY')

def configure_app():
    """Função para configurar o app"""
    st.set_page_config(
        page_title="Documentador de Power BI",
        page_icon="🧊",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    
    st.header('Documentador de Power BI')
    st.write("""Este aplicativo tem como objetivo principal facilitar a organização, o acompanhamento e a análise de dados, fornecendo uma documentação completa dos relatórios da organização. 
    Observação: a requisição para o retorno de relatórios e workspaces pode atingir o limite máximo de 50 solicitações por hora ou 15 por minuto, sendo necessário aguardar um tempo para realizar novas solicitações.""")

def sidebar_inputs():
    """Função responsável por mostrar o menu lateral e receber os dados de administrador"""
    with st.sidebar:
        st.title("Documentador de Power BI")
        st.write('Preencha com as informações do App')
        app_id = st.text_input(label='App ID')
        tenant_id = st.text_input(label='Tenant ID')
        secret_value = st.text_input(label='Secret value')
        
        st.write('')

        uploaded_files = st.file_uploader("Apenas arquivo '.pbit'", accept_multiple_files=False, type=['pbit', 'zip'], help="""Para obter o arquivo .pbit, você pode salvar o arquivo selecionando a opção para essa extensão. 
        Ou também, em Arquivo > Exportar > Power BI Template.""")
        
    return app_id, tenant_id, secret_value, uploaded_files

def main_content(headers=None, uploaded_files=None):
    """Função que mostra as informações principais do APP"""
    if uploaded_files:
        df_normalized = upload_file(uploaded_files)
        buttons_download(df_normalized)
    
    if headers:
        workspace_dict = get_workspaces_id(headers)
        
        if workspace_dict:
            option = st.selectbox("Qual workspace você gostaria de visualizar?", list(workspace_dict.keys()), index=None, placeholder='Selecione a workspace...')
            if option:
                with st.spinner('Retornando relatório...'):
                    workspace_id = workspace_dict[option]
                    scan_response = scan_workspace(headers, workspace_id)
                    display_reports(scan_response)

def display_reports(scan_response):
    """Função responsável por mostrar os paineis e lidar com a seleção"""
    report_names = [report_info['name'] for report_info in scan_response['datasets'] if 'PbixInImportMode' in report_info['contentProviderType'] and 'Usage Metrics Report' not in report_info['name']]
    
    option = st.selectbox("Qual relatório você gostaria de visualizar?", list(report_names), index=None, placeholder='Selecione o relatório...')
    
    if option:
        df_desnormalized = clean_reports(scan_response, option)
        buttons_download(df_desnormalized)

def buttons_download(df):
    on = st.toggle("Mostrar tabela completa")

    if on:
        st.dataframe(df)
    
    col1, col2, col3 = st.columns(3)

    with col1:   
        with st.spinner('Gerando tabela completa para excel...'):
            
            buffer = BytesIO()
            with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                df.to_excel(writer, sheet_name='Sheet1')
                writer.close()
            
            st.download_button(
                label="Baixar tabela completa para Excel",
                data=buffer,
                file_name=f'desnormalizada.xlsx',
                mime="application/vnd.ms-excel"
            )

            if st.button('Mostrar apenas as tabelas'):
                filtered_df = df[df['NomeTabela'].notnull() & df['FonteDados'].notnull()]
                filtered_df = filtered_df[['NomeTabela', 'FonteDados']].drop_duplicates().reset_index(drop=True)
                st.dataframe(filtered_df)

    with col2:
        if st.button("Documentar painel para Excel"):
            text, measures_df = text_to_document(df)
            buffer = generate_excel(text, measures_df)
            st.session_state['buffer'] = buffer
        
            if 'buffer' in st.session_state:
                st.download_button(
                    label="Baixar planilha",
                    data=st.session_state['buffer'],
                    file_name='painel_documentado.xlsx',
                    mime="application/vnd.ms-excel"
                )
                
        if st.button('Mostrar apenas as colunas'):
            filtered_df = df[df['NomeColuna'].notnull() & df['TipoDadoColuna'].notnull() & df['ExpressaoColuna'].notnull()]
            filtered_df = filtered_df[['NomeTabela', 'NomeColuna', 'TipoDadoColuna', 'ExpressaoColuna']].drop_duplicates().reset_index(drop=True)
            st.dataframe(filtered_df)

    with col3:
        if st.button("Documentar painel para Word"):
            text, measures_df = text_to_document(df)
            doc = generate_docx(text, measures_df)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Baixar Word",
                data=buffer,
                file_name='document.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.button('Mostrar apenas as medidas'):
            filtered_df = df[df['NomeMedida'].notnull() & df['ExpressaoMedida'].notnull()]
            filtered_df = filtered_df[['NomeMedida', 'ExpressaoMedida']].drop_duplicates().reset_index(drop=True)
            st.dataframe(filtered_df)

def text_to_document(df):
    """Texto que será inserido no prompt do bot"""    
    tables_df = df[df['NomeTabela'].notnull() & df['FonteDados'].notnull()]
    tables_df = tables_df[['NomeTabela', 'FonteDados']].drop_duplicates().reset_index(drop=True)
    
    measures_df = df[df['NomeMedida'].notnull() & df['ExpressaoMedida'].notnull()]
    measures_df = measures_df[['NomeMedida', 'ExpressaoMedida']].drop_duplicates().reset_index(drop=True)
    
    document_text = f"""
    Relatório: {df['ReportName'].iloc[0]}
    
    Tabelas:
    {tables_df['NomeTabela'].to_string(index=False)}
    
    Fontes dos dados das tabelas:
    {tables_df.to_string(index=False)}
    
    Medidas:
    {measures_df.to_string(index=False)}
    """ 
        
    return document_text, measures_df

def main():
    """Função principal do app, onde tudo será apresentado"""
    configure_app()

    app_id, tenant_id, secret_value, uploaded_files = sidebar_inputs()
    
    if app_id and tenant_id and secret_value:
        headers = get_token(app_id, tenant_id, secret_value)
        if headers:
            main_content(headers, None)
    
    if uploaded_files:
        main_content(None, uploaded_files)
        
            
def upload_file(uploaded_files):
    """Processa o upload do arquivo .pbit ou .zip e extrai os dados relevantes."""
    if uploaded_files.name.endswith('.pbit') or uploaded_files.name.endswith('.zip'):
        if uploaded_files.name.endswith('.pbit'):
            uploaded_files.name = uploaded_files.name[:-5] + '.zip'
            
        with ZipFile(uploaded_files, 'r') as zipf:
            zipf.extractall('temp')
            file_list = zipf.namelist()
            
            for file_name in file_list:
                with zipf.open(file_name) as extracted_file:
                    if file_name == 'Connections':
                        connections_content = extracted_file.read().decode("utf-8")                        
                        
                        connections_content = json.loads(connections_content)
                        
                        datasetid_content = connections_content['RemoteArtifacts'][0]['DatasetId']
                        reportid_content = connections_content['RemoteArtifacts'][0]['ReportId']
                        reportname_content = uploaded_files.name[:-4]
                    
                    if file_name == 'DataModelSchema':
                        content = extracted_file.read().decode("utf-16-le")                        
                        content = json.loads(content)
                        
        df_columns, df_tables = pd.DataFrame(), pd.DataFrame()
        
        measure_names, measure_expression, tables_names = [], [], []
        
        if 'model' in content and 'tables' in content['model']:
            tables = content['model']['tables']            
            for rows in tables:
                if 'DateTable' not in rows['name']:
                    if 'measures' in rows:
                        for measures in rows['measures']:
                            tables_names.append(rows['name'])
                            measure_names.append(measures['name'])
                            measure_expression.append("".join(measures['expression']))
                    
                    for cols in rows['columns']:
                        col_data = pd.DataFrame([{
                            'NomeTabela': rows['name'],
                            'NomeColuna': cols['name'],
                            'TipoDadoColuna': cols['dataType'],
                            'TipoColuna': cols.get('type', 'N/A'),
                            'ExpressaoColuna': cols.get('expression', 'N/A')
                        }])

                        df_columns = pd.concat([df_columns, col_data], ignore_index=True)

                    mcode = [''.join(rows['partitions'][0]['source']['expression'])]
                    
                    df_tables_rows = pd.DataFrame([{
                        'DatasetId': datasetid_content,
                        'ReportId': reportid_content,
                        'ReportName': reportname_content,
                        'NomeTabela': rows['name'], 
                        'FonteDados': mcode[0]
                    }])

                    df_tables = pd.concat([df_tables, df_tables_rows], ignore_index=True)

                    
        df_columns['ExpressaoColuna'] = df_columns['ExpressaoColuna'].apply(lambda l: "".join(l))
        
        df_measures = pd.DataFrame({
            'NomeTabela': tables_names,
            'NomeMedida': measure_names,
            'ExpressaoMedida': measure_expression
        })

        df_normalized = pd.merge(pd.merge(df_tables, df_measures, left_on='NomeTabela', right_on='NomeTabela', how='left'), df_columns, right_on='NomeTabela', left_on='NomeTabela', how='left')
        
        return df_normalized
    else:
        st.write('Arquivo não suportado')
            
def get_token(APP_ID, TENANT_ID, SECRET_VALUE):
    """Função para pegar o token do cliente da Microsoft"""
    APP_ID = APP_ID
    TENANT_ID = TENANT_ID
    SECRET_VALUE = SECRET_VALUE
    authority = f"https://login.microsoftonline.com/{TENANT_ID}"
    scopes = ["https://analysis.windows.net/powerbi/api/.default"]

    app = msal.ConfidentialClientApplication(APP_ID, authority=authority, client_credential=SECRET_VALUE)
    result = app.acquire_token_for_client(scopes=scopes)
    access_token = result["access_token"]

    headers = {
        'Authorization': f"Bearer {access_token}",
        "Content-Type": "application/json",
    }

    return headers

def get_workspaces_id(headers):
    """Função para pegar o id e as workspaces, é utilizado um handle para lidar com código 429 (Too many requests)"""
    retries = 5
    workspaces_url = 'https://api.powerbi.com/v1.0/myorg/admin/groups?$top=100'

    for i in range(retries):
        response_workspaces = requests.get(url=workspaces_url, headers=headers)
        if response_workspaces.status_code == 200:
            workspaces = response_workspaces.json().get('value', [])
            workspace_dict = {workspace['name']: workspace['id'] for workspace in workspaces}
            return workspace_dict
        elif response_workspaces.status_code == 429:
            time.sleep(2 ** i)
        else:
            st.error(f"Erro: {response_workspaces.status_code}")
            break
    return None

def scan_workspace(headers, workspace_id):
    """Função responsável por fazer um escaneamento na workspace e recuperar suas informações.
    Utiliza dados da função get_workspaces_id para passar a workspaceid no body"""
    url = 'https://api.powerbi.com/v1.0/myorg/admin/workspaces/getInfo?datasetSchema=True&datasetExpressions=True'
    body = {
    "workspaces": [
        f'{workspace_id}'
    ]
    }

    response = requests.post(url=url, headers=headers, json=body)
    scanId = response.json()['id']
        
    time.sleep(5)
    # Para esperar o scan finalizar
    
    scan_url = f'https://api.powerbi.com/v1.0/myorg/admin/workspaces/scanResult/{scanId}'

    scan_response = requests.get(url=scan_url, headers=headers)
    reports = scan_response.json()['workspaces'][0]
    
    return reports

def clean_reports(reports, option):
    """Função responsável por fazer a limpeza do JSON que é recebido através da API da Microsoft, ao serem inseridos as credenciais do APP, e logo após o armazena-lo em um Pandas DataFrame"""

    df_workspaces = pd.json_normalize(reports).explode('datasets', ignore_index=True)

    # Filtrando e tratando o dataset principal
    df_normalized = pd.json_normalize(df_workspaces['datasets'])
    df_normalized = df_normalized.query(f"name == '{option}'")
    df_normalized = df_normalized[['id', 'name', 'configuredBy', 'createdDate', 'tables', 'expressions']].copy()
    df_normalized.rename(columns={'id': 'DatasetId', 'name': 'ReportName'}, inplace=True)
    datasets_exploded = df_normalized.explode('tables', ignore_index=True)

    # Normalizando e tratando a tabela de 'tables'
    tables_normalized = pd.concat([datasets_exploded[['DatasetId', 'ReportName', 'configuredBy']], pd.json_normalize(datasets_exploded['tables'])], axis=1)
    tables_normalized.rename(columns={'name': 'NomeTabela'}, inplace=True)
    tables_normalized['source'] = tables_normalized['source'].apply(lambda x: x[0]['expression'] if isinstance(x, list) and len(x) > 0 else None)
    if 'storageMode' not in tables_normalized.columns:
        tables_normalized['storageMode'] = None

    # Criando e tratando a tabela de medidas
    measures_normalized = tables_normalized.explode('measures', ignore_index=True)
    measures_normalized = pd.concat([measures_normalized[['NomeTabela']], pd.json_normalize(measures_normalized['measures'])], axis=1)
    measures_normalized['name'] = measures_normalized.get('name', 'N/A')
    measures_normalized['expression'] = measures_normalized.get('expression', 'N/A')
    measures_normalized = measures_normalized[['NomeTabela', 'name', 'expression']]
    measures_normalized.rename(columns={'name': 'NomeMedida', 'expression': 'ExpressaoMedida'}, inplace=True)

    # Criando e tratando a tabela de colunas
    columns_normalized = tables_normalized.explode('columns', ignore_index=True)
    columns_normalized = pd.concat([columns_normalized[['NomeTabela']], pd.json_normalize(columns_normalized['columns'])], axis=1)
    columns_normalized = columns_normalized[['NomeTabela', 'name', 'dataType', 'columnType', 'expression']]
    columns_normalized['expression'] = columns_normalized.get('expression', 'N/A')
    columns_normalized.rename(columns={'name': 'NomeColuna', 'dataType': 'TipoDadoColuna', 'columnType': 'TipoColuna', 'expression': 'ExpressaoColuna'}, inplace=True)

    tables_normalized = tables_normalized[['DatasetId', 'ReportName', 'NomeTabela', 'storageMode', 'source', 'configuredBy']]
    tables_normalized.rename(columns={'source': 'FonteDados'}, inplace=True)
    dataset_desnormalized = tables_normalized.merge(measures_normalized, on='NomeTabela', how='left')
    dataset_desnormalized = dataset_desnormalized.merge(columns_normalized, on='NomeTabela', how='left')

    return dataset_desnormalized

def prompt():
    prompt_relatorio = """
    Você é um documentador especializado em Power BI. Sua função é criar documentações claras e detalhadas para os relatórios, tabelas, medidas e fontes de dados em Power BI. Para cada item, você deve incluir uma descrição compreensiva que ajude os usuários a entenderem sua finalidade e uso no contexto do relatório. Utilize uma linguagem técnica e precisa, mas acessível para usuários com diferentes níveis de conhecimento em Power BI.
    Fazer a documentação em JSON.
    Você deverá dividir em diferentes outupt de acordo com a entrada do usuário, separando em: info_paineis, tabelas, medidas e fonte_de_dados.
    Para a parte de medidas, você deverá fazer em blocos, das que estiverem sendo solicitadas, mas como continuação do JSON e ao final de todas fechar o JSON igual no exemplo.
    Retorne apenas o json, sem o ```json no inicio e o ``` no final

    Instruções Específicas:

    Relatórios:
    - Título do Relatório
    - Descrição do objetivo do relatório
    - Principais KPIs e métricas apresentadas
    - Público-alvo do relatório
    - Exemplos de uso

    Formato de Documentação:

    Tabelas do Relatório
    Nome da Tabela | Descrição da Tabela

    Medidas do Relatório
    Nome da Medida | Descrição da Medida

    Fontes de Dados
    Nome da Fonte de Dados | Descrição da Fonte | Tabelas Contidas no M

    Exemplo de Documentação:

    {
    "Relatorio": {
        "Titulo": "Análise de Vendas Mensais",
        "Descricao": "Este relatório fornece uma visão detalhada das vendas mensais por região e produto. Os principais KPIs incluem receita total, unidades vendidas e margem de lucro. O relatório é destinado aos gerentes de vendas regionais e é atualizado semanalmente para refletir os dados mais recentes.",
        "Principais_KPIs_e_Metricas": [
        "Receita Total",
        "Unidades Vendidas",
        "Margem de Lucro"
        ],
        "Publico_Alvo": "Gerentes de Vendas Regionais",
        "Exemplos_de_Uso": [
        "Identificação de tendências de vendas por região",
        "Comparação de desempenho de produtos"
        ]
    },
    "Tabelas_do_Relatorio": [
        {
        "Nome": "Vendas",
        "Descricao": "Tabela que armazena dados de vendas, incluindo ID do produto, quantidade vendida, preço e data da venda."
        },
        {
        "Nome": "Produtos",
        "Descricao": "Tabela que contém informações detalhadas dos produtos, como nome, categoria e preço unitário."
        }
    ],
    "Medidas_do_Relatorio": [
        {
        "Nome": "Receita Total",
        "Descricao": "Calcula a receita total das vendas somando o preço de venda multiplicado pela quantidade vendida."
        },
        {
        "Nome": "Margem de Lucro",
        "Descricao": "Calcula a margem de lucro subtraindo o custo do preço de venda."
        }
    ],
    "Fontes_de_Dados": [
        {
        "Nome": "SQL Server - Vendas",
        "Descricao": "Base de dados contendo todas as transações de vendas da empresa.",
        "Tabelas_Contidas_no_M": [
            "Vendas",
            "Produtos"
        ]
        },
        {
        "Nome": "Excel - Orçamento",
        "Descricao": "Planilha contendo dados de orçamento anual por departamento.",
        "Tabelas_Contidas_no_M": [
            "Orçamento"
        ]
        }
    ]
    }

    Abaixo estão dados do relatório do Power BI:
    <INICIO DADOS RELATORIO POWER BI>

    """
    return prompt_relatorio

def client_chat(messages):
    client = OpenAI(api_key=API_KEY)
    
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        max_tokens=4096,
        messages=messages
    )

    return json.loads(response.choices[0].message.content)

def Documenta(prompt, text):
    messages = [
        {"role": "system", "content": "Você é um documentador especializado em Power BI."},
        {"role": "user", "content": f"{prompt}\n{text}\n<FIM DADOS RELATORIO POWER BI>"}
    ]

    messages.append({"role": "user", "content": "Para essa solicitação você deverá apenas retornar a parte do json 'Relatorio'"})
    response_info = client_chat(messages)
    
    messages.append({"role": "user", "content": "Para essa solicitação você deverá apenas retornar a parte do json 'Tabelas_do_Relatorio'"})
    response_tables = client_chat(messages)
    
    messages.append({"role": "user", "content": "Para essa solicitação você deverá apenas retornar a parte do json 'Medidas_do_Relatorio'. Se a medida for NaN, não retorne ela."})
    response_measures = client_chat(messages)
    
    messages.append({"role": "user", "content": "Para essa solicitação você deverá apenas retornar a parte do json 'Fontes_de_Dados'"})
    response_source = client_chat(messages)
    
    return response_info, response_tables, response_measures, response_source

def generate_docx(text, measures_df):
    doc = Document()
 
    prompts = prompt()
 
    response_info, response_tables, response_measures, response_source = Documenta(prompts, text)
    
    doc.add_paragraph(f'Título do relatório: {response_info["Relatorio"]["Titulo"]}')
    doc.add_paragraph(f'Descrição: {response_info["Relatorio"]["Descricao"]}')
    doc.add_paragraph(f'Principais KPIs e Métricas: {", ".join(response_info["Relatorio"]["Principais_KPIs_e_Metricas"])}')
    doc.add_paragraph(f'Público alvo: {response_info["Relatorio"]["Publico_Alvo"]}')
    doc.add_paragraph(f'Exemplos de uso: {", ".join(response_info["Relatorio"]["Exemplos_de_Uso"])}\n')

    doc.add_paragraph('Tabelas do relatório\n')
    
    if 'Tabelas_do_Relatorio' in response_tables:
        for table in response_tables['Tabelas_do_Relatorio']:
            doc.add_paragraph(f'Tabela: {table["Nome"]}\nDescrição: {table["Descricao"]}\n')
    else:
        for table in response_tables:
            doc.add_paragraph(f'Tabela: {table["Nome"]}\nDescrição: {table["Descricao"]}\n')

    doc.add_paragraph('Medidas do relatório\n')
    
    # Para não fazer o chat repetir a expressão ela é pega por um dataframe
    def add_measure_paragraph():
        measure_name = measure["Nome"]
        expression = measures_df.loc[measures_df['NomeMedida'] == measure_name, 'ExpressaoMedida'].values[0]
        doc.add_paragraph(f'Nome: {measure_name}\nDescrição: {measure["Descricao"]}\nFórmula DAX: {expression}\n')

    if 'Medidas_do_Relatorio' in response_measures and 'Nome' in response_measures['Medidas_do_Relatorio']:
        for measure in response_measures['Medidas_do_Relatorio']:
            add_measure_paragraph(measure)
            
    elif isinstance(response_tables, list):
        for measure in response_measures:
            add_measure_paragraph(measure)
    else:
        doc.add_paragraph('O relatório não possui medidas\n')

    doc.add_paragraph('Fonte de dados do relatório\n')
    
    if 'Fontes_de_Dados' in response_source:
        for source in response_source['Fontes_de_Dados']:
            doc.add_paragraph(f'Nome: {source["Nome"]}\nDescrição: {source["Descricao"]}\nTabelas contidas no M: {", ".join(source["Tabelas_Contidas_no_M"])}\n')
    else:
        for source in response_source:
            doc.add_paragraph(f'Nome: {source["Nome"]}\nDescrição: {source["Descricao"]}\nTabelas contidas no M: {", ".join(source["Tabelas_Contidas_no_M"])}\n')

    return doc

def generate_excel(text, measures_df):
    """Função responsável por tratar e gerar o excel do output do chatgpt"""
    prompts = prompt()
    buffer = BytesIO()
    
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        all_info = []
        all_tables = []
        all_measures = []
        all_sources = []
        
        response_info, response_tables, response_measures, response_source = Documenta(prompts, text)
        
        info = pd.DataFrame([response_info['Relatorio']]).transpose()
        info.reset_index(inplace=True)
        info.columns = ['Informações do relatório', 'Dados']
        all_info.append(info)
        
        tables = pd.DataFrame(response_tables['Tabelas_do_Relatorio'])
        all_tables.append(tables)
        
        measures = pd.DataFrame(response_measures['Medidas_do_Relatorio'])
        all_measures.append(measures)
        
        sources = pd.DataFrame(response_source['Fontes_de_Dados'])
        all_sources.append(sources)
    
        df_info = pd.concat(all_info, ignore_index=True)
        df_tabelas = pd.concat(all_tables, ignore_index=True)
        df_medidas = pd.concat(all_measures, ignore_index=True)
        df_fontes = pd.concat(all_sources, ignore_index=True)
        
        if 'Nome' in response_measures['Medidas_do_Relatorio']:
            df_medidas = pd.merge(df_medidas, measures_df,  left_on='Nome', right_on='Medida', how='left')
            df_medidas = df_medidas[['Medida', 'Descricao', 'expression']]
    
        df_info.to_excel(writer, sheet_name='info_painel', index=False)
        df_tabelas.to_excel(writer, sheet_name='tabelas', index=False) 
        df_medidas.to_excel(writer, sheet_name='medidas', index=False) 
        df_fontes.to_excel(writer, sheet_name='fonte_de_dados', index=False) 
            
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    main()
